import docx

def parse_docx(file_path: str) -> str:
    """
    Extracts text from paragraphs and table cells in a DOCX file using python-docx.
    """
    doc = docx.Document(file_path)
    paragraphs = []
    
    # Extract text from standard paragraphs
    for p in doc.paragraphs:
        paragraphs.append(p.text)
        
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Add paragraph text from tables
                for p in cell.paragraphs:
                    paragraphs.append(p.text)
                    
    # Join and return non-empty paragraphs
    return "\n".join(p for p in paragraphs if p.strip())
